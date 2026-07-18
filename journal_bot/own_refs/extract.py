"""PDF → Plain-Text → Refs-Section → DOIs + Free-Text-Citations.

Portiert aus `scripts/iter11_extract_own_refs.py` (Iter-11-Backtest, dort
gegen 109 echte PDFs validiert) mit folgenden Änderungen für den Produktiv-
Track:
- Cache-Verzeichnis ist konfigurierbar (Default:
  `<project_root>/.own_refs_cache/fulltext/`).
- Cache-Key ist sha1(absolute pdf-path), nicht zotero_key — Folder-Sources
  haben keinen Zotero-Key.
- Funktionen liefern Dataclass-Strukturen statt JSON-Files zu schreiben.
- Idempotenz: ist der `.txt`-Cache jünger als das PDF, wird pdftotext
  übersprungen.

KEINE LLM-Calls. Reine pdftotext + Regex.
"""

from __future__ import annotations

import hashlib
import re
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_FULLTEXT_CACHE = PROJECT_ROOT / ".own_refs_cache" / "fulltext"

# Auf macOS Homebrew. Auf Linux i. d. R. /usr/bin/pdftotext.
PDFTOTEXT_CANDIDATES = (
    "/opt/homebrew/bin/pdftotext",
    "/usr/local/bin/pdftotext",
    "/usr/bin/pdftotext",
    "pdftotext",
)


# -- Regexes ------------------------------------------------------------------

HEADER_RE = re.compile(
    r"^[ \t]*("
    r"references"
    r"|literatur(?:verzeichnis)?"
    r"|bibliograph(?:y|ie)"
    r"|bibliografie"
    r"|works[ \t]+cited"
    r"|quellen(?:verzeichnis)?"
    r"|cited[ \t]+works"
    r"|literatur[ \t]+und[ \t]+quellen"
    r")[ \t:.]*$",
    re.IGNORECASE,
)

# Erweiterte Header-Erkennung (§2.3). Akzeptiert:
#  - Section-Prefix: "4." / "VIII." / "II."  vor Header
#  - Header-Suffix:  "Literaturverzeichnis I.1; I.4; I.5"   (Sammelband)
#  - Sammelband-Sub-Headers: "Primärliteratur", "Sekundärliteratur"
#  - OCR-Artefakt: "rEFErENcEs", "Bibliogr aphy" (Spacing-Glitch)
# WICHTIG: in find_references_block werden Inhalts-Verzeichnis-Einträge
# (Header + Seitenzahl am Ende) explizit ausgefiltert.
HEADER_RE_EXTENDED = re.compile(
    r"^[ \t]*"
    r"(?:(?:\d+|[IVXivx]+)\.[ \t]+)?"          # optionaler Section-Prefix
    r"("
    r"references"
    r"|literatur(?:verzeichnis)?"
    r"|bibliograph(?:y|ie)"
    r"|bibliografie"
    r"|works[ \t]+cited"
    r"|quellen(?:verzeichnis)?"
    r"|cited[ \t]+works"
    r"|literatur[ \t]+und[ \t]+quellen"
    r"|prim(?:ä|ae)rliteratur"                  # Sammelband
    r"|sekund(?:ä|ae)rliteratur"                # Sammelband
    r"|verwendete[ \t]+literatur"
    r"|zitierte[ \t]+literatur"
    r"|sources(?:[ \t]+cited)?"
    r")"
    r"(?:[ \t]+[IVXivx0-9.;,\- ]+)?"            # optionaler Header-Suffix
    r"[ \t:.]*$",
    re.IGNORECASE,
)

# Inhaltsverzeichnis-Erkennung: Header durch ≥3 Spaces oder Punkt-Linie
# getrennt von einer Seitenzahl am Ende. Heuristik bewusst eng — Sammelband-
# Section-References wie "Literaturverzeichnis I.1; I.4; I.5" sind KEIN TOC.
TOC_ENTRY_RE = re.compile(
    r"(?:\s{3,}|\.{3,}\s*|…\s*)\d{1,4}\s*$"
)

DOI_RE = re.compile(r"\b10\.\d{4,9}/[A-Za-z0-9._;()/:\-]+", re.IGNORECASE)

CITATION_START_RE = re.compile(
    r"^[A-ZÄÖÜ][A-Za-zÄÖÜäöüß\-'\.]+(?:,\s+|[ \t]+(?:and|und|&)\s+)"
)

POST_REFS_CUT_RE = re.compile(
    r"^[ \t]*("
    r"(?:open[ \t]+)?access\s+this\s+chapter\s+is\s+licensed"
    r"|über[ \t]+(?:die[ \t]+)?autor"
    r"|about[ \t]+the[ \t]+author"
    r"|biographische[ \t]+notiz"
    r"|autorenangaben"
    r"|impressum"
    r"|copyright[ \t]+©"
    r"|appendix"
    r"|anhang"
    r")\b",
    re.IGNORECASE,
)


# -- Dataclasses --------------------------------------------------------------


@dataclass
class ExtractionResult:
    pdf_path: str
    txt_path: str | None
    fulltext_chars: int = 0
    refs_text: str = ""
    refs_header_line: int | None = None
    refs_header_label: str | None = None
    used_fallback_section: bool = False
    dois_in_refs: list[str] = field(default_factory=list)
    raw_citations: list[str] = field(default_factory=list)
    status: str = "ok"                  # ok | no_pdf | pdftotext_failed | pdf_empty
    notes: list[str] = field(default_factory=list)


# -- pdftotext ----------------------------------------------------------------


def _resolve_pdftotext() -> str | None:
    import shutil
    for cand in PDFTOTEXT_CANDIDATES:
        if Path(cand).exists() or shutil.which(cand):
            return cand
    return None


def _pdf_cache_path(pdf_path: Path, cache_dir: Path) -> Path:
    h = hashlib.sha1(str(pdf_path.resolve()).encode("utf-8")).hexdigest()[:16]
    return cache_dir / f"{h}.txt"


# -- Offene Formate -----------------------------------------------------------
#
# Für die eigenen Publikationen liegen meist Manuskripte in offenen Formaten
# vor. Die sind der PDF-Extraktion vorzuziehen: pdftotext `-layout` verschränkt
# bei zweispaltigem Satz die Spalten zeilenweise (gemessen 2026-07-18: 12 von 73
# Volltexten mit ≥30 % Spaltenrinne, Spitze 83 %), und gescannte PDFs tragen
# zusätzlich OCR-Zeichenfehler, die kein Schalter behebt. Auf Seitenzahlen und
# Schlussredaktion kommt es für die Werkanalyse nicht an.

OPEN_TEXT_SUFFIXES = frozenset({".txt", ".md", ".markdown"})
OPEN_DOC_SUFFIXES = frozenset({".docx", ".odt"})
SUPPORTED_SUFFIXES = frozenset({".pdf"}) | OPEN_TEXT_SUFFIXES | OPEN_DOC_SUFFIXES


def _text_from_docx(path: Path) -> str:
    """Absätze und Tabellenzellen eines .docx in Lesereihenfolge."""
    import docx  # python-docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def _text_from_odt(path: Path) -> str:
    """Text eines .odt ohne Zusatzabhängigkeit: content.xml entpacken, Tags weg.

    `text:p`/`text:h` werden zu Zeilenumbrüchen, damit Absatzgrenzen erhalten
    bleiben — die Refs-Heuristiken arbeiten zeilenweise.
    """
    import xml.etree.ElementTree as ET
    import zipfile

    with zipfile.ZipFile(path) as z:
        xml = z.read("content.xml").decode("utf-8", errors="replace")
    root = ET.fromstring(xml)
    ns = "{urn:oasis:names:tc:opendocument:xmlns:text:1.0}"
    lines: list[str] = []
    for el in root.iter():
        if el.tag in (f"{ns}p", f"{ns}h"):
            lines.append("".join(el.itertext()))
    return "\n".join(lines)


def ensure_fulltext(
    pdf_path: Path,
    cache_dir: Path = DEFAULT_FULLTEXT_CACHE,
    force: bool = False,
) -> tuple[Path | None, str]:
    """Stellt sicher, dass für `pdf_path` ein Plain-Text-Cache existiert.

    Nimmt neben PDF auch offene Formate (.docx, .odt, .txt, .md) — für die
    eigenen Publikationen sind die der PDF-Extraktion vorzuziehen (siehe
    OPEN_DOC_SUFFIXES).

    Returns: (txt_path | None, status). Status ∈ {ok, no_pdf, pdftotext_failed,
    pdftotext_missing, open_format_failed}.
    Idempotent: wenn txt-Cache aktueller als die Quelldatei, wird nicht erneut
    extrahiert.
    """
    if not pdf_path.exists() or not pdf_path.is_file():
        return None, "no_pdf"

    cache_dir.mkdir(parents=True, exist_ok=True)
    txt_path = _pdf_cache_path(pdf_path, cache_dir)

    if not force and txt_path.exists():
        try:
            if txt_path.stat().st_mtime >= pdf_path.stat().st_mtime:
                return txt_path, "ok"
        except OSError:
            pass

    suffix = pdf_path.suffix.lower()
    if suffix in OPEN_TEXT_SUFFIXES or suffix in OPEN_DOC_SUFFIXES:
        try:
            if suffix in OPEN_TEXT_SUFFIXES:
                text = pdf_path.read_text(encoding="utf-8", errors="replace")
            elif suffix == ".docx":
                text = _text_from_docx(pdf_path)
            else:
                text = _text_from_odt(pdf_path)
        except Exception as e:  # noqa: BLE001 — Formatfehler nicht vorab klassifizierbar
            print(f"  [warn] {suffix} extraction failed for {pdf_path}: {e}", file=sys.stderr)
            return None, "open_format_failed"
        txt_path.write_text(text, encoding="utf-8")
        return txt_path, "ok"

    pdftotext = _resolve_pdftotext()
    if pdftotext is None:
        return None, "pdftotext_missing"

    try:
        subprocess.run(
            [pdftotext, "-layout", "-enc", "UTF-8", str(pdf_path), str(txt_path)],
            check=True, capture_output=True, timeout=120,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
        print(f"  [warn] pdftotext failed for {pdf_path}: {e}", file=sys.stderr)
        return None, "pdftotext_failed"

    return (txt_path, "ok") if txt_path.exists() else (None, "pdftotext_failed")


# -- Refs-Section + Heuristiken ----------------------------------------------


def _is_toc_entry(line: str) -> bool:
    """Heuristik: Header + Seitenzahl am Ende = TOC-Eintrag, kein Refs-Header.
    'VIII. Literaturverzeichnis  249' → True.
    """
    return bool(TOC_ENTRY_RE.search(line))


def find_references_block(
    text: str,
    *,
    use_fallback: bool = True,
) -> tuple[str, int | None, str | None]:
    """Liefert (refs_text, header_line_no, header_label).

    Strategie (§2.3):
    1. Suche das LETZTE Header-Vorkommen ab 30 % Dokumentenlänge mit der
       erweiterten Regex (HEADER_RE_EXTENDED). Filtert TOC-Einträge raus.
    2. Bei Sammelbänden mit "Primärliteratur" + "Sekundärliteratur":
       nimm den ERSTEN Sub-Header ab 30 %, damit beide Sektionen
       eingeschlossen sind.
    3. Wenn `use_fallback=True` und kein Header gefunden wurde: nutze die
       letzten 25 % als Pseudo-Refs (markiert mit header_label="(fallback)").
       Das fängt Spalten-Layouts und Editorial-Notizen ab.
    """
    lines = text.splitlines()
    n = len(lines)
    if n == 0:
        return "", None, None
    min_line = int(n * 0.30)
    header_line = None
    header_label = None
    first_sammelband_line = None
    first_sammelband_label = None

    for i, line in enumerate(lines):
        if i < min_line:
            continue
        stripped = line.strip()
        m = HEADER_RE_EXTENDED.match(stripped)
        if not m:
            continue
        if _is_toc_entry(stripped):
            continue
        label = m.group(1).lower()
        # Sammelband-Sub-Header: ersten Treffer merken (umfasst beide Sektionen)
        if label.startswith(("prim", "sekund")) and first_sammelband_line is None:
            first_sammelband_line = i
            first_sammelband_label = m.group(1)
        # Normaler Header: nimm den letzten
        header_line = i
        header_label = m.group(1)

    # Wenn Sammelband-Header gefunden, prefer den ersten Sub-Header
    if first_sammelband_line is not None:
        header_line = first_sammelband_line
        header_label = first_sammelband_label

    if header_line is not None:
        return "\n".join(lines[header_line + 1:]), header_line, header_label

    # Fallback: letzte 25 % als Pseudo-Refs-Sektion, markiert.
    # Greift bei Sammelbänden ohne klare Header, Editorials, Spalten-Layouts.
    if use_fallback and n >= 40:
        fallback_start = int(n * 0.75)
        return "\n".join(lines[fallback_start:]), fallback_start, "(fallback)"

    return "", None, None


def cut_post_refs_garbage(refs_text: str) -> str:
    """Schneidet typische Anhänge nach der Refs-Section ab."""
    lines = refs_text.splitlines()
    for i, line in enumerate(lines):
        if POST_REFS_CUT_RE.match(line):
            return "\n".join(lines[:i])
    return refs_text


def extract_dois(text: str) -> list[str]:
    raw = DOI_RE.findall(text)
    seen: set[str] = set()
    out: list[str] = []
    for d in raw:
        d = d.rstrip(".,;:)]")
        d = re.sub(r"-\s*\n\s*", "", d)
        d_lc = d.lower()
        if d_lc not in seen and len(d_lc) > 7:
            seen.add(d_lc)
            out.append(d_lc)
    return out


def split_citations(refs_text: str) -> list[str]:
    """Heuristisches Splitting in einzelne Citations.

    Eine neue Citation startet, wenn die Zeile links beginnt (≤ 3 Spaces)
    und mit `Lastname, …` / `Lastname and …` / `Lastname & …` einsetzt.
    Folgezeilen werden eingerückt erwartet und gehören zur laufenden Citation.
    """
    lines = refs_text.splitlines()
    citations: list[list[str]] = []
    current: list[str] = []
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            if current:
                citations.append(current)
                current = []
            continue
        leading = len(line) - len(line.lstrip(" \t"))
        is_start = leading <= 3 and bool(CITATION_START_RE.match(line.lstrip()))
        if is_start and current:
            citations.append(current)
            current = [line.strip()]
        else:
            current.append(line.strip())
    if current:
        citations.append(current)
    out: list[str] = []
    for block in citations:
        joined = re.sub(r"\s+", " ", " ".join(block)).strip()
        if len(joined) >= 20:
            out.append(joined)
    return out


# -- Hauptfunktion ------------------------------------------------------------


def extract_refs(
    pdf_path: Path,
    cache_dir: Path = DEFAULT_FULLTEXT_CACHE,
    force: bool = False,
) -> ExtractionResult:
    """Vollständige Extraktion eines PDFs: Fulltext → Refs-Section → DOIs + Citations."""
    txt_path, status = ensure_fulltext(pdf_path, cache_dir, force=force)
    if txt_path is None:
        return ExtractionResult(
            pdf_path=str(pdf_path), txt_path=None, status=status,
            notes=[f"ensure_fulltext={status}"],
        )

    text = txt_path.read_text(encoding="utf-8", errors="replace")
    chars = len(text)
    if not text.strip():
        return ExtractionResult(
            pdf_path=str(pdf_path), txt_path=str(txt_path),
            fulltext_chars=chars, status="pdf_empty",
        )

    refs_text, header_line, header_label = find_references_block(text)
    # find_references_block kann drei Zustände liefern:
    #   - Header gefunden:           label = "Literatur"/"References"/etc.
    #   - Letzte-25%-Fallback:       label = "(fallback)"  (§2.3)
    #   - Gar nichts (kurzes Doc):   label = None, refs_text = ""
    # In den letzten zwei Fällen ist Citation-Splitting sinnlos, weil
    # die "Sektion" Fließtext sein kann.
    used_fallback = header_label in (None, "(fallback)")
    if header_label is None and not refs_text:
        # Notfall-Fallback: ganzer Volltext (alter Pre-§2.3-Pfad, fängt
        # sehr kurze Dokumente ohne Header-Match ab).
        refs_text = text

    refs_text = cut_post_refs_garbage(refs_text)
    dois = extract_dois(refs_text)
    citations = [] if used_fallback else split_citations(refs_text)

    return ExtractionResult(
        pdf_path=str(pdf_path),
        txt_path=str(txt_path),
        fulltext_chars=chars,
        refs_text=refs_text,
        refs_header_line=header_line,
        refs_header_label=header_label,
        used_fallback_section=used_fallback,
        dois_in_refs=dois,
        raw_citations=citations,
        status="ok",
    )
